import os
import base64
import csv
import io
import json
import re
import urllib.request
from datetime import datetime
import streamlit as st
from dotenv import load_dotenv
from google import genai
from google.genai import types
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

load_dotenv()

NAME_COLUMN_CANDIDATES = ["氏名", "名前", "職員名", "スタッフ名", "name", "Name"]
MEETING_COLUMN_CANDIDATES = ["会議名", "会議", "名称", "meeting", "Meeting", "name", "Name"]


def fetch_column_values(csv_url: str, column_candidates: list[str]) -> list[str]:
    """Googleスプレッドシートの「ウェブに公開」CSVリンクから、見出し行が候補に一致する列の値を取得する"""
    req = urllib.request.Request(csv_url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(req, timeout=15) as resp:
        raw = resp.read().decode("utf-8-sig")

    reader = csv.DictReader(io.StringIO(raw))
    fieldnames = reader.fieldnames or []
    target_field = next(
        (f for f in fieldnames if f.strip() in column_candidates),
        fieldnames[0] if fieldnames else None,
    )

    values = []
    if target_field:
        for row in reader:
            value = (row.get(target_field) or "").strip()
            if value and value not in values:
                values.append(value)
    return values


def fetch_attendee_roster(csv_url: str) -> list[str]:
    return fetch_column_values(csv_url, NAME_COLUMN_CANDIDATES)


def fetch_meeting_list(csv_url: str) -> list[str]:
    return fetch_column_values(csv_url, MEETING_COLUMN_CANDIDATES)


def sanitize_filename_part(text: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|]', "_", text).strip()
    return cleaned or "議事録"


DRIVE_SCOPES = ["https://www.googleapis.com/auth/drive"]


def get_drive_service():
    credentials_b64 = os.environ.get("GOOGLE_DRIVE_CREDENTIALS_JSON_B64")
    if credentials_b64:
        # Streamlit Community Cloudなど、JSONキーファイルを配置できない環境向け
        # （改行や記号を含む生JSONは貼り付け時に壊れやすいため、Base64化した1行の文字列をst.secretsに保存する）
        info = json.loads(base64.b64decode(credentials_b64).decode("utf-8"))
        creds = service_account.Credentials.from_service_account_info(info, scopes=DRIVE_SCOPES)
    else:
        creds = service_account.Credentials.from_service_account_file(
            os.environ["GOOGLE_DRIVE_CREDENTIALS_FILE"], scopes=DRIVE_SCOPES
        )
    return build("drive", "v3", credentials=creds)


def get_drive_root_folder_ids() -> list[str]:
    """カテゴリ別ルートフォルダ（病院／患者別／看護部／病棟・部署など）の共有ドライブ上のフォルダID一覧を.envから取得する"""
    raw = os.environ.get("GOOGLE_DRIVE_ROOT_FOLDER_IDS") or ""
    return [folder_id.strip() for folder_id in raw.split(",") if folder_id.strip()]


def find_or_create_meeting_folder_drive(service, root_folder_ids: list[str], meeting_name: str) -> tuple[str, bool]:
    """複数のカテゴリフォルダを横断して会議名と同じ名前のフォルダを探す。
    どこにも見つからなければ、先頭のカテゴリフォルダ直下に新規作成する。
    戻り値は (フォルダID, 新規作成したかどうか)。"""
    safe_name = sanitize_filename_part(meeting_name).replace("'", "\\'")
    for root_id in root_folder_ids:
        query = (
            f"'{root_id}' in parents and name = '{safe_name}' "
            "and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        )
        response = service.files().list(
            q=query,
            fields="files(id, name)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
        ).execute()
        matches = response.get("files", [])
        if matches:
            return matches[0]["id"], False

    created = service.files().create(
        body={
            "name": safe_name,
            "mimeType": "application/vnd.google-apps.folder",
            "parents": [root_folder_ids[0]],
        },
        fields="id",
        supportsAllDrives=True,
    ).execute()
    return created["id"], True


def save_minutes_to_drive(local_docx_path: str, meeting_name: str, filename: str) -> tuple[str | None, bool]:
    """議事録ファイルを、会議名と同じ名前のフォルダ（複数のカテゴリフォルダを横断検索）にGoogle Drive API経由でアップロードする。
    戻り値は (共有ドライブ上のファイルURL, 新規フォルダを作成したかどうか)。未設定の場合は (None, False)。"""
    root_folder_ids = get_drive_root_folder_ids()
    if not root_folder_ids:
        return None, False

    service = get_drive_service()
    folder_id, created_new_folder = find_or_create_meeting_folder_drive(service, root_folder_ids, meeting_name)

    media = MediaFileUpload(
        local_docx_path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    uploaded = service.files().create(
        body={"name": filename, "parents": [folder_id]},
        media_body=media,
        fields="webViewLink",
        supportsAllDrives=True,
    ).execute()
    return uploaded.get("webViewLink"), created_new_folder

# ==========================================
# 1. 出力JSONスキーマの定義
# ==========================================
class Attendee(BaseModel):
    name: str
    role: str
    present: bool = Field(description="出席チェック(出点)")

class ActionItem(BaseModel):
    owner: str
    task: str
    due: str

class AgendaCoverage(BaseModel):
    agenda: str
    decided: bool
    note: str

class MinutesResponse(BaseModel):
    meeting_type: str
    date: str
    attendees: list[Attendee]
    summary: str = Field(description="会議全体の要約 (3~5文)")
    decisions: list[str]
    action_items: list[ActionItem]
    agenda_coverage: list[AgendaCoverage]
    unresolved: list[str]

# ==========================================
# 2. Gemini API 呼び出し処理 (★参加者リストを受け取るように変更)
# ==========================================
def generate_minutes(audio_file_path: str, meeting_template: str, terms_list: str, attendees_list: list) -> str:
    client = genai.Client()
    audio_file = client.files.upload(file=audio_file_path)

    # リストをカンマ区切りの文字列に変換
    attendees_str = "、".join(attendees_list) if attendees_list else "指定なし"

    system_instruction = f"""
    あなたは医療機関のプロフェッショナルな議事録作成AIです。
    提供された音声データから、以下の会議テンプレートと専門用語集に従って構造化された議事録を作成してください。

    【参加予定者名簿】
    {attendees_str}
    ※音声の内容とこの名簿を照合し、参加が確認できた者は present: true、欠席の場合は present: false としてください。

    【会議テンプレート】
    {meeting_template}

    【専門用語集】
    {terms_list}
    """

    response = client.models.generate_content(
        model='gemini-3.5-flash',
        contents=[audio_file, "この会議音声を元に議事録を生成してください。"],
        config=types.GenerateContentConfig(
            system_instruction=system_instruction,
            response_mime_type="application/json",
            response_schema=MinutesResponse,
            temperature=0.2,
        ),
    )
    return response.text

# ==========================================
# 3. Wordファイル(.docx)作成処理
# ==========================================
def create_docx_from_json(json_data: str, output_filename: str = "minutes_output.docx"):
    data = json.loads(json_data)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"{data.get('meeting_type', '会議')} 議事録")
    title_run.font.size = Pt(18)
    title_run.bold = True
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sign_table = doc.add_table(rows=2, cols=3)
    sign_table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_table.style = 'Table Grid'

    headers = ["事務長", "課長(確認)", "作成者"]
    for i, h in enumerate(headers):
        cell = sign_table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in sign_table.rows[1:]:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(r'<w:trHeight %s w:val="600" w:hRule="atLeast"/>' % nsdecls('w')))

    doc.add_paragraph()

    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'

    info_table.cell(0, 0).text = "開催日時"
    info_table.cell(0, 1).text = f"{data.get('date', '未定')}"

    info_table.cell(1, 0).text = "開催場所"
    info_table.cell(1, 1).text = "院内会議室"

    attendees = [f"{a['name']}" for a in data.get('attendees', []) if a.get('present')]
    info_table.cell(2, 0).text = "出席者"
    info_table.cell(2, 1).text = "、".join(attendees) if attendees else "記載なし"

    for row in info_table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.0)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    doc.add_heading("1. 会議要約", level=1)
    doc.add_paragraph(data.get('summary', ''))

    doc.add_heading("2. 決定事項", level=1)
    for decision in data.get('decisions', []):
        doc.add_paragraph(decision, style='List Bullet')

    doc.add_heading("3. ネクストアクション（担当・期限）", level=1)
    for item in data.get('action_items', []):
        doc.add_paragraph(f"【{item['owner']}】 {item['task']} (期限: {item['due']})", style='List Bullet')

    doc.add_heading("4. 保留・次回持ち越し事項", level=1)
    for unresolved in data.get('unresolved', []):
        doc.add_paragraph(unresolved, style='List Bullet')

    doc.save(output_filename)
    return output_filename

# ==========================================
# 4. Streamlit Webアプリ画面の構築
# ==========================================
st.set_page_config(page_title="議事録AI自動化システム", page_icon="🎙️", layout="wide")

CUSTOM_CSS = """
<style>
:root {
    --primary: #0F6E68;
    --primary-dark: #0B4F4A;
    --border: #E3E8E8;
}
html, body, [class*="css"] {
    font-family: -apple-system, BlinkMacSystemFont, "Hiragino Sans",
        "Hiragino Kaku Gothic ProN", "Yu Gothic", Meiryo, sans-serif;
}
.block-container {
    max-width: 1120px;
    padding-top: 2rem;
    padding-bottom: 3rem;
}
.hero {
    background: linear-gradient(135deg, var(--primary-dark), var(--primary));
    border-radius: 16px;
    padding: 2.2rem 2.5rem;
    color: #fff;
    margin-bottom: 1.8rem;
    box-shadow: 0 8px 24px rgba(15, 110, 104, 0.18);
}
.hero h1 {
    font-size: 1.9rem;
    margin: 0 0 0.4rem 0;
    font-weight: 700;
}
.hero p {
    margin: 0;
    font-size: 0.98rem;
    opacity: 0.92;
}
.step-label {
    display: inline-block;
    background: var(--primary);
    color: #fff;
    font-size: 0.72rem;
    font-weight: 700;
    letter-spacing: 0.04em;
    padding: 0.2rem 0.6rem;
    border-radius: 999px;
    margin-bottom: 0.6rem;
}
div[data-testid="stVerticalBlockBorderWrapper"] {
    border-radius: 14px !important;
}
.attendee-chip {
    display: inline-flex;
    align-items: center;
    gap: 0.35rem;
    background: #EEF5F4;
    color: var(--primary-dark);
    border: 1px solid #D6E7E5;
    border-radius: 999px;
    padding: 0.3rem 0.8rem;
    margin: 0.2rem 0.35rem 0.2rem 0;
    font-size: 0.85rem;
    font-weight: 600;
}
.attendee-chip.absent {
    background: #F5F5F5;
    color: #9AA3A3;
    border-color: #E5E5E5;
    text-decoration: line-through;
}
.stButton>button, .stDownloadButton>button {
    border-radius: 10px;
    font-weight: 700;
    padding: 0.6rem 1.4rem;
}
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

st.markdown(
    """
    <div class="hero">
        <h1>🎙️ 京阪病院 議事録生成AI</h1>
        <p>会議の音声をアップロードすると、AIが自動で病院フォーマットの議事録（Wordファイル）を作成します。</p>
    </div>
    """,
    unsafe_allow_html=True,
)

if not os.environ.get("GEMINI_API_KEY"):
    st.error(
        "GEMINI_API_KEYが設定されていません。アプリと同じフォルダの `.env` ファイルに\n"
        "`GEMINI_API_KEY=あなたのAPIキー` の形式で1行追加してから、アプリを再起動してください。"
    )
    st.stop()

TEMPLATE = "1. 再発防止策の決定\n2. 前回対策の進捗確認"
TERMS = "精神科特有の語彙（例: 弄便, 連合弛緩, マル障）"

col1, col2, col3 = st.columns(3, gap="large")

with col1:
    with st.container(border=True):
        st.markdown('<span class="step-label">STEP 1</span>', unsafe_allow_html=True)
        st.markdown("#### 🗂️ 会議名の選択")

        meeting_url = os.environ.get("MEETING_LIST_CSV_URL")

        if "meeting_list" not in st.session_state:
            if meeting_url:
                try:
                    fetched_meetings = fetch_meeting_list(meeting_url)
                    st.session_state.meeting_list = fetched_meetings
                    if fetched_meetings:
                        st.session_state.meeting_list_updated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
                except Exception as e:
                    st.session_state.meeting_list = []
                    st.session_state.meeting_list_fetch_error = str(e)
            else:
                st.session_state.meeting_list = []

        meeting_refresh_clicked = st.button(
            "🔄 スプレッドシートから再取得",
            key="meeting_refresh",
            use_container_width=True,
            disabled=not meeting_url,
            help=None if meeting_url else "MEETING_LIST_CSV_URLが.envに未設定です",
        )
        if meeting_refresh_clicked and meeting_url:
            try:
                fetched_meetings = fetch_meeting_list(meeting_url)
                if fetched_meetings:
                    st.session_state.meeting_list = fetched_meetings
                    st.session_state.meeting_list_updated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
                    st.success(f"会議名一覧を更新しました（{len(fetched_meetings)}件）")
                else:
                    st.warning("スプレッドシートから会議名を取得できませんでした。列名や内容をご確認ください。")
            except Exception as e:
                st.error(f"会議名一覧の取得に失敗しました: {e}")

        if st.session_state.get("meeting_list_updated_at"):
            st.caption(f"最終更新: {st.session_state.meeting_list_updated_at}（スプレッドシートより取得）")
        elif st.session_state.get("meeting_list_fetch_error"):
            st.caption(f"スプレッドシートの取得に失敗しました（{st.session_state.meeting_list_fetch_error}）。")
        elif meeting_url:
            st.caption("スプレッドシートと連携しています。")
        else:
            st.caption("MEETING_LIST_CSV_URLが.envに未設定です。")

        selected_meeting = st.selectbox(
            "会議名を選択してください",
            options=st.session_state.meeting_list,
            index=None,
            placeholder="会議名を選択…",
            label_visibility="collapsed",
        )

with col2:
    with st.container(border=True):
        st.markdown('<span class="step-label">STEP 2</span>', unsafe_allow_html=True)
        st.markdown("#### 👥 参加予定者の設定")

        roster_url = os.environ.get("ATTENDEE_ROSTER_CSV_URL")

        if "attendee_roster" not in st.session_state:
            if roster_url:
                try:
                    fetched = fetch_attendee_roster(roster_url)
                    st.session_state.attendee_roster = fetched
                    if fetched:
                        st.session_state.roster_updated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
                except Exception as e:
                    st.session_state.attendee_roster = []
                    st.session_state.roster_fetch_error = str(e)
            else:
                st.session_state.attendee_roster = []

        refresh_clicked = st.button(
            "🔄 スプレッドシートから再取得",
            key="roster_refresh",
            use_container_width=True,
            disabled=not roster_url,
            help=None if roster_url else "ATTENDEE_ROSTER_CSV_URLが.envに未設定です",
        )
        if refresh_clicked and roster_url:
            try:
                fetched = fetch_attendee_roster(roster_url)
                if fetched:
                    st.session_state.attendee_roster = fetched
                    st.session_state.roster_updated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
                    st.success(f"名簿を更新しました（{len(fetched)}名）")
                else:
                    st.warning("スプレッドシートから名前を取得できませんでした。列名や内容をご確認ください。")
            except Exception as e:
                st.error(f"名簿の取得に失敗しました: {e}")

        if st.session_state.get("roster_updated_at"):
            st.caption(f"最終更新: {st.session_state.roster_updated_at}（スプレッドシートより取得）")
        elif st.session_state.get("roster_fetch_error"):
            st.caption(f"スプレッドシートの取得に失敗しました（{st.session_state.roster_fetch_error}）。")
        elif roster_url:
            st.caption("スプレッドシートと連携しています。")
        else:
            st.caption("ATTENDEE_ROSTER_CSV_URLが.envに未設定です。")

        current_roster = st.session_state.attendee_roster
        default_selection = []
        selected_attendees = st.multiselect(
            "参加予定者を選択してください（複数選択可）",
            options=current_roster,
            default=default_selection,
            label_visibility="collapsed",
        )

with col3:
    with st.container(border=True):
        st.markdown('<span class="step-label">STEP 3</span>', unsafe_allow_html=True)
        st.markdown("#### 📁 音声データのアップロード")
        uploaded_file = st.file_uploader(
            "会議の音声ファイルを選択してください (m4a, mp3, wav)",
            type=['m4a', 'mp3', 'wav'],
            label_visibility="collapsed",
        )
        if uploaded_file is not None:
            st.success(f"「{uploaded_file.name}」を読み込みました")

with st.container(border=True):
    st.markdown('<span class="step-label">STEP 4</span>', unsafe_allow_html=True)
    st.markdown("#### ✨ 議事録の生成")
    generate_clicked = st.button(
        "議事録を生成する",
        type="primary",
        disabled=uploaded_file is None or not selected_meeting,
        use_container_width=True,
    )
    if not selected_meeting:
        st.caption("先に会議名を選択してください。")
    elif uploaded_file is None:
        st.caption("先に音声ファイルをアップロードしてください。")

if uploaded_file is not None and selected_meeting and generate_clicked:
    temp_audio_path = "temp_uploaded_audio" + os.path.splitext(uploaded_file.name)[1]
    try:
        with st.status("議事録を作成しています…", expanded=True) as status:
            st.write("① 音声ファイルを準備中…")
            with open(temp_audio_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            st.write("② AIが音声を解析中…（内容により数十秒〜数分かかります）")
            result_json = generate_minutes(temp_audio_path, TEMPLATE, TERMS, selected_attendees)
            data = json.loads(result_json)
            data["meeting_type"] = selected_meeting
            result_json = json.dumps(data, ensure_ascii=False)

            st.write("③ 病院フォーマットのWordを作成中…")
            file_base = f"{sanitize_filename_part(selected_meeting)}_{datetime.now().strftime('%Y%m%d_%H%M')}"
            docx_filename = create_docx_from_json(result_json, output_filename=f"{file_base}.docx")

            saved_path = None
            save_error = None
            saved_new_folder = False
            drive_save_configured = bool(get_drive_root_folder_ids())
            if drive_save_configured:
                st.write("④ 会議名フォルダへアップロード中…")
                try:
                    saved_path, saved_new_folder = save_minutes_to_drive(
                        docx_filename, selected_meeting, f"{file_base}.docx"
                    )
                except Exception as e:
                    save_error = str(e)

            status.update(label="議事録の作成が完了しました", state="complete", expanded=False)

        with st.container(border=True):
            st.markdown("#### 📝 議事録プレビュー")

            if drive_save_configured:
                if saved_path:
                    st.success(f"「{selected_meeting}」フォルダにアップロードしました → {saved_path}")
                    if saved_new_folder:
                        st.caption(f"「{selected_meeting}」という名前のフォルダが見つからなかったため、新規作成しました。")
                elif save_error:
                    st.warning(f"Google Driveへの自動アップロードに失敗しました（下のボタンからダウンロードできます）: {save_error}")

            st.markdown(f"##### {data.get('meeting_type', '会議')}")
            meta_col1, meta_col2 = st.columns(2)
            meta_col1.metric("開催日", data.get("date", "-"))
            present_count = sum(1 for a in data.get("attendees", []) if a.get("present"))
            meta_col2.metric("出席者数", f"{present_count}名")

            chips_html = "".join(
                f'<span class="attendee-chip{"" if a.get("present") else " absent"}">'
                f'{"✓" if a.get("present") else "✕"} {a.get("name", "")}</span>'
                for a in data.get("attendees", [])
            )
            st.markdown(chips_html, unsafe_allow_html=True)

            st.write("")
            tab_summary, tab_decisions, tab_actions, tab_agenda, tab_unresolved = st.tabs(
                ["要約", "決定事項", "ネクストアクション", "議題の消化状況", "保留事項"]
            )

            with tab_summary:
                st.write(data.get("summary", ""))

            with tab_decisions:
                decisions = data.get("decisions", [])
                if decisions:
                    for d in decisions:
                        st.markdown(f"- {d}")
                else:
                    st.caption("決定事項はありません。")

            with tab_actions:
                items = data.get("action_items", [])
                if items:
                    for item in items:
                        st.markdown(f"**【{item['owner']}】** {item['task']}　（期限: {item['due']}）")
                else:
                    st.caption("ネクストアクションはありません。")

            with tab_agenda:
                agenda_items = data.get("agenda_coverage", [])
                if agenda_items:
                    for ag in agenda_items:
                        icon = "✅" if ag.get("decided") else "🕒"
                        st.markdown(f"{icon} **{ag.get('agenda', '')}**")
                        if ag.get("note"):
                            st.caption(ag["note"])
                else:
                    st.caption("議題の消化状況の記録はありません。")

            with tab_unresolved:
                unresolved_items = data.get("unresolved", [])
                if unresolved_items:
                    for u in unresolved_items:
                        st.markdown(f"- {u}")
                else:
                    st.caption("保留・次回持ち越し事項はありません。")

            st.write("")
            with open(docx_filename, "rb") as file:
                st.download_button(
                    label="📥 病院フォーマットWordをダウンロード",
                    data=file,
                    file_name=f"{file_base}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
    except Exception as e:
        st.error(f"エラーが発生しました: {e}")
    finally:
        if os.path.exists(temp_audio_path):
            os.remove(temp_audio_path)